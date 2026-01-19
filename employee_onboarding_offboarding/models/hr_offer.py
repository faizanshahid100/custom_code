from email.policy import default
from odoo import models, fields, api
from odoo.exceptions import ValidationError, UserError
import base64
from docx import Document
import os
import subprocess
from datetime import date, timedelta
from dateutil.relativedelta import relativedelta



class HrOffer(models.Model):
    _name = "hr.offer"
    _description = "Employee Offer Creation"
    _inherit = ["mail.thread", "mail.activity.mixin"]

    # Candidate Details
    candidate_name = fields.Char("Candidate Name", required=True, tracking=True)
    # contact_number = fields.Char("Contact Number", required=True, tracking=True)
    personal_email = fields.Char("Personal Email", required=True, tracking=True)
    official_email = fields.Char("Official Email", tracking=True)
    gazetted_holiday_id = fields.Many2one('gazetted.holiday', string='Gazetted Holiday Policy', tracking=True)
    employee_id = fields.Many2one('hr.employee', string='Employee', ondelete='cascade')
    country_id = fields.Many2one('res.country', string="Country", required=True, tracking=True)
    joining_date = fields.Date("Joining Date", required=True, tracking=True)
    job_id = fields.Many2one("hr.job", string="Designation", required=True, tracking=True)
    currency_id = fields.Many2one('res.currency', string='Currency', required=True,
                                  default=lambda self: self.env.company.currency_id)
    salary = fields.Float("Salary", tracking=True)
    allowances = fields.Float("Allowances", tracking=True)
    employment_type = fields.Selection([
        ("permanent", "Permanent"),
        ("contract", "Contract"),
        ("intern", "Internship"),
    ], string="Employment Type", required=True, tracking=True)
    probation_period = fields.Integer("Probation Period (months)", default=3, tracking=True)

    # Client Assignment
    client_id = fields.Many2one("res.partner", string="Client", required=True, domain="[('is_company','=',True)]", tracking=True)
    client_joining_date = fields.Date("Client Joining Date", tracking=True)
    pillar = fields.Selection([
        ("business", "Business"),
        ("tech", "Tech"),
        ("management", "Management"),
    ], string="Business/Tech Pillar", required=True, tracking=True)
    work_location = fields.Selection([
        ("onsite", "On-Site"),
        ("fully_remote", "Remote"),
        ("hybrid", "Hybrid"),
    ], string="Work Location",default='onsite', required=True, tracking=True)

    # TODO remove this field and related model in Odoo 19 while migrating the code
    onsite_day_ids = fields.Many2many('hr.onsite.day', string="Onsite Days")
    reporting_time = fields.Char("Reporting Time", default='12:00 PM - 9:00 PM (PST PK)', tracking=True)
    working_hours = fields.Float("Working Hours (per day)", default=9.0, tracking=True)
    checklist_template_id = fields.Many2one('checklist.template', string='Checklist Template', tracking=True)
    contract_type = fields.Selection([
        ("pakistan", "Pakistan"),
        ("philippines", "Philippines"),
    ], string='Contract Type', required=True, tracking=True)
    hr_responsible = fields.Selection([
        ("pakistan", "Pakistan"),
        ("philippines", "Philippines"),
    ], string='Hr Responsible', required=True, tracking=True)

    # Reporting Structure
    manager_id = fields.Many2one("hr.employee", string="Internal Manager", required=True, tracking=True)
    manager = fields.Char(string='Manager (Client)')
    manager_email = fields.Char(string='Manager Email')
    department_id = fields.Many2one('hr.department', string='Department')
    dept_hod = fields.Char(string="Dept HOD", tracking=True)
    dept_hod_email = fields.Char(string="HOD Email", readonly=True, tracking=True)

    # Additional Notes
    buddy_info = fields.Char("Buddy Info", tracking=True)
    buddy_id = fields.Many2one('hr.employee', string='Buddy Name', tracking=True)
    remarks = fields.Text("Modification Remarks", tracking=True)
    special_instructions = fields.Text("Special Instructions", tracking=True)
    internal_remarks = fields.Text("Internal Remarks", tracking=True)
    offer_submitter_id = fields.Many2one('res.users', string='Offer Submitter')

    # Workflow
    state = fields.Selection([
        ("draft", "New"),
        ("submitted", "CEO Approval"),
        ("modification", "Modification"),
        ("approved", "Approved"),
        ("send_offer", "Offer Sent"),
        ("send_contract", "Contract Sent"),
        ("hired", "Hired"),
        ("rejected", "Rejected"),
    ], default="draft", tracking=True)

    # Contract Info (Existing fields also be used here)
    date_of_birth = fields.Date(string='Date of Birth', tracking=True)
    address = fields.Char(string='Address', tracking=True)
    candidate_mobile = fields.Char(string='Candidate Mobile', tracking=True)
    id_number = fields.Char(string='ID number/Passport', tracking=True)
    tax_number = fields.Char(string='Tax ID (NTN, TIN)', tracking=True)
    ice_name = fields.Char(string='ICE Name', tracking=True)
    ice_number = fields.Char(string='ICE Number', tracking=True)
    ice_relation = fields.Char(string='ICE Relation', tracking=True)
    bank_name = fields.Char(string='Bank Name', tracking=True)
    bank_info = fields.Char(string='Branch City and Branch Code', tracking=True)
    account_number = fields.Char(string='Account Number', tracking=True)
    iban_number = fields.Char(string='IBAN Number', tracking=True)
    swift_code = fields.Char(string='SWIFT Code', tracking=True)
    linked_in_profile = fields.Char(string='LinkedIn Profile Link', tracking=True)

    # Job Description
    job_description = fields.Html(string='Job Description', tracking=True)

    def name_get(self):
        result = []
        for rec in self:
            result.append((rec.id, '%s - ( %s)' % (rec.candidate_name, rec.client_id.name)))
        return result

    # Workflow Actions
    def action_submit(self):
        self.write({
            "state": "submitted",
            "offer_submitter_id":self.env.user.id
        })

        # Get CEO group users
        ceo_group = self.env.ref("employee_onboarding_offboarding.group_ceo")
        ceo_users = ceo_group.users

        if not ceo_users:
            return

        # Get mail template
        template = self.env.ref("employee_onboarding_offboarding.candidate_offer_approval_template")

        if template:
            for user in ceo_users:
                # Send mail to each CEO
                template.send_mail(self.id, force_send=True, email_values={"email_to": user.email})

    def action_send_back(self):
        return {
            "type": "ir.actions.act_window",
            "res_model": "hr.offer.sendback.wizard",
            "view_mode": "form",
            "target": "new",
            "context": {"active_id": self.id},
        }

    def action_approve(self):
        if not self.env.user.has_group("employee_onboarding_offboarding.group_ceo"):
            raise ValidationError("Only CEO can approve offers.")

        self.write({"state": "approved"})

        # Find all HR Responsible users
        hr_users = self.env['res.users'].search([
            ('groups_id', 'in', self.env.ref("employee_onboarding_offboarding.group_responsible_hr").id)
        ])

        # Send email notification
        if hr_users:
            for user in hr_users:
                if user.login:
                    mail_values = {
                        'subject': "Offer Approved by CEO",
                        'body_html': f"<p>Dear {user.name},</p>"
                                     f"<p>The CEO has approved the offer <b>{self.candidate_name}</b>. "
                                     f"Now you can proceed to send it to the employee.</p>",
                        'email_to': user.email,
                    }
                    self.env['mail.mail'].sudo().create(mail_values).send()

    def action_sent_offer(self):
        self.ensure_one()
        if not self.env.user.has_group("employee_onboarding_offboarding.group_responsible_hr"):
            raise ValidationError("Only HR Responsible can send offers.")

        self.write({"state": "send_contract"})

        if not self.personal_email:
            raise ValidationError("Personal email is required to send offer.")

        if self.contract_type == 'pakistan':
            template = self.env.ref(
                "employee_onboarding_offboarding.candidate_offer_final_template",
                raise_if_not_found=False
            )
        elif self.contract_type == 'philippines':
            template = self.env.ref(
                "employee_onboarding_offboarding.candidate_offer_final_template_philippines",
                raise_if_not_found=False
            )
        else:
            raise ValidationError("Unsupported contract type.")

        template.send_mail(self.id, force_send=True)

    def action_sent_contract(self):
        for record in self:
            record.write({"state": "send_contract"})

            # Get module base path
            module_path = os.path.dirname(os.path.abspath(__file__))
            base_path = os.path.join(module_path, "..")

            # Template in /data/
            # template_path = os.path.join(base_path, "data", "contract_template.docx")
            # TODO : when below done then above line to be removed
            if self.contract_type == 'pakistan':
                template_path = os.path.join(base_path, "data", "contract_template.docx")
            elif self.contract_type == 'philippines':
                template_path = os.path.join(base_path, "data", "contract_template_philippines.docx")

            # Contracts folder
            contracts_dir = os.path.join(base_path, "contracts")
            if not os.path.exists(contracts_dir):
                os.makedirs(contracts_dir)

            # Load template
            doc = Document(template_path)

            # Replace variables
            replacements = {
                "{{ candidate_name }}": record.candidate_name or "",
                "{{ candidate_designation }}": record.job_id.name or "",
                "{{ id_number }}": record.id_number or "",
                "{{ address }}": record.address or "",
                "{{ offer_issue_date }}": fields.Date.today().strftime("%d %B %Y"),
                "{{ job_description }}": record.job_description or "",
                "{{ joining_date }}": record.joining_date.strftime("%d %B %Y") if record.joining_date else "",
                "{{ probation_end_date }}": ((record.joining_date + relativedelta(months=3 if record.contract_type == 'pakistan' else 6 )).strftime("%d %B %Y") if record.joining_date else ""),
                "{{ salary }}": str(record.salary) if record.salary else "",
                "{{ allowances }}": str(record.allowances) if record.allowances else "0",
                "{{ special_instructions }}": str(record.special_instructions) if record.special_instructions else "None",
                "{{ currency }}": str(record.currency_id.name),
                "{{ probation_period }}": str(record.probation_period or ""),
                "{{ work_location }}": record.work_location + ' ' + (', '.join(record.onsite_day_ids.mapped('name'))) if record.work_location == 'hybrid' else record.work_location,
                "{{ reporting_time }}": record.reporting_time or "",
            }

            for para in doc.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)

            # Save final doc in contracts folder
            safe_name = record.candidate_name.replace(" ", "_")+'_'+record.id_number
            final_path = os.path.join(contracts_dir, f"Contract_{safe_name}.docx")
            doc.save(final_path)

            # Attach to Odoo
            with open(final_path, "rb") as f:
                file_data = f.read()

            # Convert docx to PDF using LibreOffice
            pdf_path = os.path.join(contracts_dir, f"Contract_{safe_name}.pdf")
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', contracts_dir, final_path], check=True)
            
            # Read PDF file
            with open(pdf_path, "rb") as f:
                pdf_data = f.read()

            attachment = self.env["ir.attachment"].create({
                "name": f"Contract_{safe_name}.pdf",
                "type": "binary",
                "datas": base64.b64encode(pdf_data),
                "res_model": "hr.offer",
                "res_id": record.id,
                "mimetype": "application/pdf",
            })

            # --- SEND EMAIL WITH ATTACHMENT ---
            hr_responsible_group = (
                'employee_onboarding_offboarding.group_responsible_hr_pak'
                if record.contract_type == 'pakistan'
                else 'employee_onboarding_offboarding.group_responsible_hr_philippines'
            )

            mail_values = {
                "subject": "Independent Contractor Service Agreement",
                "body_html": f"""
                                <p>Dear <b>{record.candidate_name}</b>,</p>
                                <p>Please find attached your Independent Contractor Service Agreement.</p>
                                <p>Kindly review, sign, and reply back with confirmation.</p>
                                <br/>
                                <p><b>Regards,</b><br/>
                                HR Team<br/>
                                Prime System Solutions</p>
                            """,
                "email_from": "hr@primesystemsolutions.com",
                "email_to": record.personal_email,
                "email_cc": "misbah.yasir@primesystemsolutions.com" if record.contract_type == 'pakistan' else "misbah.yasir@primesystemsolutions.com,sharo.domingo@primesystemsolutions.com,patricia.reyes@primesystemsolutions.com",
                "attachment_ids": [(6, 0, [attachment.id])],
            }
            self.env["mail.mail"].sudo().create(mail_values).send()

    def action_hire(self):
        for record in self:
            # Change state
            record.write({"state": "hired"})

            # Validations
            if not record.gazetted_holiday_id:
                raise ValidationError('Please enter Gazetted Holiday first')
            elif not record.checklist_template_id:
                raise ValidationError('Please enter Checklist Template first')

            if record.employment_type != 'intern':
                if not record.official_email:
                    raise ValidationError('Please enter Official Email first')


                user = self.env['res.users'].sudo().create({
                    'name': record.candidate_name,
                    'login': record.official_email,
                    'password': '12345',
                })

                # Create Employee
                employee_vals = {
                    "name": record.candidate_name,
                    "work_email": record.official_email,
                    "work_phone": record.candidate_mobile,
                    "mobile_phone": record.candidate_mobile,
                    "private_email": record.personal_email,
                    "birthday": record.date_of_birth,
                    # "address_home_id": record.address_id.id if record.address_id else False,
                    "job_id": record.job_id.id if record.job_id else False,
                    "department_id": record.department_id.id if record.department_id else False,
                    "parent_id": record.manager_id.id if record.manager_id else False,
                    "coach_id": record.manager_id.id if record.manager_id else False,
                    "country_id": record.country_id.id if record.country_id else False,
                    # "bank_account_id": record.bank_account_id.id if hasattr(record, 'bank_account_id') else False,
                    # "work_location_id": record.work_location_id.id if hasattr(record, 'work_location_id') else False,
                    "manager": record.manager,
                    "manager_email": record.manager_email,
                    "dept_hod": record.dept_hod,
                    "dept_hod_email": record.dept_hod_email,
                    "joining_date": record.joining_date,
                    "joining_salary": record.salary,
                    "contractor": record.client_id.id,
                    "work_mode": record.work_location,
                    "onsite_day_ids": record.onsite_day_ids.ids,
                    "total_working_hour": record.working_hours,
                    "gazetted_holiday_id": record.gazetted_holiday_id.id,
                    "checklist_template_id": record.checklist_template_id.id,
                    "emergency_contact": record.ice_name,
                    "emergency_phone": record.ice_number,
                    "emergency_contact_relation": record.ice_relation,
                    "identification_id": record.id_number,
                    "user_id": user.id,
                }

                employee = self.env["hr.employee"].sudo().create(employee_vals)

                # (Optional) link back offer → employee
                record.write({"employee_id": employee.id})

    def action_make_profile(self):
        for record in self:
            # Validations
            if not record.gazetted_holiday_id:
                raise ValidationError('Please enter Gazetted Holiday first')
            elif not record.checklist_template_id:
                raise ValidationError('Please enter Checklist Template first')

            if not record.official_email:
                raise ValidationError('Please enter Official Email first')

            user = self.env['res.users'].sudo().create({
                'name': record.candidate_name,
                'login': record.official_email,
                'password': '12345',
            })

            # Create Employee
            employee_vals = {
                "name": record.candidate_name,
                "work_email": record.official_email,
                "work_phone": record.candidate_mobile,
                "mobile_phone": record.candidate_mobile,
                "private_email": record.personal_email,
                "birthday": record.date_of_birth,
                # "address_home_id": record.address_id.id if record.address_id else False,
                "job_id": record.job_id.id if record.job_id else False,
                "department_id": record.department_id.id if record.department_id else False,
                "parent_id": record.manager_id.id if record.manager_id else False,
                "coach_id": record.manager_id.id if record.manager_id else False,
                "country_id": record.country_id.id if record.country_id else False,
                # "bank_account_id": record.bank_account_id.id if hasattr(record, 'bank_account_id') else False,
                # "work_location_id": record.work_location_id.id if hasattr(record, 'work_location_id') else False,
                "manager": record.manager,
                "manager_email": record.manager_email,
                "dept_hod": record.dept_hod,
                "dept_hod_email": record.dept_hod_email,
                "joining_date": record.joining_date,
                "joining_salary": record.salary,
                "work_mode": record.work_location,
                "onsite_day_ids": record.onsite_day_ids.ids,
                "total_working_hour": record.working_hours,
                "gazetted_holiday_id": record.gazetted_holiday_id.id,
                "checklist_template_id": record.checklist_template_id.id,
                "emergency_contact": record.ice_name,
                "emergency_phone": record.ice_number,
                "emergency_contact_relation": record.ice_relation,
                "identification_id": record.id_number,
                "user_id": user.id,
            }

            employee = self.env["hr.employee"].sudo().create(employee_vals)

            # (Optional) link back offer → employee
            record.write({
                "employee_id": employee.id,
                "employment_type" : "permanent"
                            })

    def action_reject(self):
        self.write({"state": "rejected"})

    def request_official_email(self):
        for record in self:
            # Get IT Support users
            it_support_group = self.env.ref("employee_onboarding_offboarding.group_it_support")
            it_support_emails = [user.email for user in it_support_group.users if user.email]

            if not it_support_emails:
                raise UserError("No IT Support user with email found.")

            # Get HR Responsible users
            hr_group = self.env.ref("employee_onboarding_offboarding.group_responsible_hr")
            hr_emails = [user.email for user in hr_group.users if user.email]

            # Build subject & body
            subject = f"Request to Create Official Email - {record.candidate_name}"
            body = f"""
                Dear IT Support Team,<br/><br/>
                Please create an official email ID for the candidate <b>{record.candidate_name}</b>.<br/>
                Personal Email: {record.personal_email}<br/><br/>
                Once created, kindly inform HR Responsible ({', '.join(hr_emails)}).<br/><br/>
                Regards,<br/>
                HR Team
            """

            # Send mail
            mail_values = {
                "subject": subject,
                "body_html": body,
                "email_from": "hr@primesystemsolutions.com",
                "email_to": ",".join(it_support_emails),
                "email_cc": ",".join(hr_emails),
            }
            self.env["mail.mail"].sudo().create(mail_values).send()

    def _cron_send_login_credentials(self):
        tomorrow = date.today() + timedelta(days=1)
        records = self.search([('joining_date', '=', tomorrow), ('official_email', '!=', False)])
        if not records:
            return

        # Get PDF file from module data folder
        module_path = os.path.dirname(os.path.abspath(__file__))
        base_path = os.path.join(module_path, "..")
        pdf_path = os.path.join(base_path, "data", "odoo_guideline_handbook.pdf")
        pdf_data = False
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_data = base64.b64encode(f.read())

        for rec in records:
            subject = "Welcome to Prime System Solutions - Your Odoo Credentials"
            body = f"""
                <p>Dear <b>{rec.candidate_name}</b>,</p>
                <p>Welcome aboard! Your professional account has been created. Below are your Odoo login credentials:</p>
                <ul>
                    <li><b>Click to Login:</b> 
                        <a href="https://eportal.primesystemsolutions.com/web?db=PrimeSystemSolution#cids=1&action=menu">
                            ePortal - Prime System Solutions
                        </a>
                    </li>
                    <li><b>User ID:</b> {rec.official_email}</li>
                    <li><b>Password:</b> 12345</li>
                </ul>
                
                <p>Welcome aboard! Your professional MS 365 account has been created. Below are your Odoo login credentials:</p>
                <ul>
                    <li><b>Click to Login:</b> 
                        <a href="https://m365.cloud.microsoft/apps/?auth=2&origindomain=microsoft365">
                            LogIn Microsoft 365
                        </a>
                    </li>
                    <li><b>User ID:</b> {rec.official_email}</li>
                    <li><b>Password:</b> Prime@123###</li>
                </ul>
                <p>For your convenience, we have attached the <b>Odoo Guideline Handbook</b> which explains how to use Odoo and change your password.</p>
                <br/>
                <p>We are excited to have you on the team!</p>
                <p>Regards,<br/>HR Team</p>
            """

            mail_values = {
                "subject": subject,
                "body_html": body,
                "email_from": "hr@primesystemsolutions.com",
                "email_to": rec.personal_email,
                "email_cc": ','.join(user.email for user in self.env.ref('employee_onboarding_offboarding.group_responsible_hr').users if user.email),
                "attachment_ids": [],
            }

            # Attach PDF
            if pdf_data:
                attachment = self.env['ir.attachment'].create({
                    'name': 'Odoo Guideline Handbook.pdf',
                    'type': 'binary',
                    'datas': pdf_data,
                    'res_model': 'hr.offer',
                    'res_id': rec.id,
                    'mimetype': 'application/pdf'
                })
                mail_values["attachment_ids"] = [(6, 0, [attachment.id])]

            self.env['mail.mail'].sudo().create(mail_values).send()

    @api.model
    def _cron_notify_buddy_before_joining(self):
        tomorrow = date.today() + timedelta(days=1)
        offers = self.search([
            ('joining_date', '=', tomorrow),
            ('buddy_id', '!=', False),
            ('state', '=', 'hired'),
        ])

        for offer in offers:
            buddy = offer.buddy_id
            if buddy.work_email:
                subject = f"New Joiner Tomorrow: {offer.candidate_name}"
                body = f"""
                        <p>Hi {buddy.name},</p>
                        <p>I’m reaching out to inform you that you’ve been assigned as the work buddy for <b>{offer.candidate_name}</b>, who joined us on <b> {offer.joining_date.strftime('%d %B %Y')} </b> as a <b> {offer.job_id.name or 'N/A'} </b>.</p>
                        <p>Your support will play a key role in helping him get settled, understand our processes, and integrate smoothly into the team. Please connect with him to offer guidance, answer any questions he may have, and share your insights about our work culture and day-to-day practices.</p>
                        </br>
                        <p>With your experience and knowledge, I’m confident that <b>{offer.candidate_name}</b> will feel welcomed and well-supported. Let me know if you need anything from my side.</p>
                        <p>Thank you for taking on this important role!</p>
                        <br/>
                        <p>Regards,<br/>HR Team</p>
                    """

                hr_users = self.env.ref('employee_onboarding_offboarding.group_responsible_hr').users
                hr_emails = ",".join([u.email for u in hr_users if u.email])
                self.env['mail.mail'].sudo().create({
                    'subject': subject,
                    'body_html': body,
                    'email_from': 'hr@primesystemsolutions.com',
                    'email_to': buddy.work_email,
                    'email_cc': hr_emails+','+offer.personal_email,  # optional: also CC candidate
                }).send()

    @api.model
    def _cron_notify_accountant_after_one_day(self):
        from datetime import date, timedelta

        yesterday = date.today() - timedelta(days=1)

        offers = self.search([
            ('joining_date', '=', yesterday),
            ('personal_email', '!=', False),
            ('state', '=', 'hired')
        ])

        if not offers:
            return

        accountant_group = self.env.ref('employee_onboarding_offboarding.group_accountant_id')
        hr_group = self.env.ref('employee_onboarding_offboarding.group_responsible_hr')

        accountant_emails = [u.email for u in accountant_group.users if u.email]
        hr_emails = [u.email for u in hr_group.users if u.email]

        if not accountant_emails:
            return

        for offer in offers:
            employee_code = offer.employee_id.barcode or ''
            salary_after_probation = (offer.salary or 0) + (offer.allowances or 0)

            # build table
            table_html = f"""
                <table border="1" cellpadding="5" cellspacing="0" style="border-collapse:collapse;">
                    <tr><th>Employee Code</th><td>{employee_code}</td></tr>
                    <tr><th>Employee CNIC</th><td>{offer.id_number or ''}</td></tr>
                    <tr><th>Employee Name</th><td>{offer.candidate_name or ''}</td></tr>
                    <tr><th>Bank Name</th><td>{offer.bank_name or ''}</td></tr>
                    <tr><th>Branch City and Branch Code</th><td>{offer.bank_info or ''}</td></tr>
                    <tr><th>Designation</th><td>{offer.job_id.name or ''}</td></tr>
                    <tr><th>Department</th><td>{offer.department_id.name or ''}</td></tr>
                    <tr><th>Joining Date</th><td>{offer.joining_date.strftime('%d %B %Y') if offer.joining_date else ''}</td></tr>
                    <tr><th>Basic Salary</th><td>{offer.salary or 0:.0f}</td></tr>
                    <tr><th>Allowances</th><td>{offer.allowances or 0:.0f}</td></tr>
                    <tr><th>Salary</th><td>{(offer.salary or 0) + (offer.allowances or 0):.0f}</td></tr>
                    <tr><th>Probation Period</th><td>{offer.probation_period or 0} months</td></tr>
                    <tr><th>Salary After Probation</th><td>{salary_after_probation:.0f}</td></tr>
                    <tr><th>Account Number</th><td>{offer.account_number or ''}</td></tr>
                    <tr><th>IBAN Number</th><td>{offer.iban_number or ''}</td></tr>
                    <tr><th>SWIFT Code</th><td>{offer.swift_code or ''}</td></tr>
                    <tr><th>Address</th><td>{offer.address or ''}</td></tr>
                </table>
                """

            subject = f"New Joiner Details - {offer.candidate_name} (Joined Yesterday)"
            body = f"""
                    <p>Dear Accounts Team,</p>
                    <p>Please find below the details of the newly joined employee:</p>
                    {table_html}
                    <br/>
                    <p>Regards,<br/>HR Team</p>
                """

            self.env['mail.mail'].sudo().create({
                'subject': subject,
                'body_html': body,
                'email_from': 'hr@primesystemsolutions.com',
                'email_to': ",".join(accountant_emails),
                'email_cc': ",".join(hr_emails + [offer.personal_email]),
            }).send()